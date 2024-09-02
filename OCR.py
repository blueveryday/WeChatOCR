import wcocr
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from colorama import init, Fore, Style

def find_wechat_path():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    common_paths = os.path.join(script_dir, 'path')
    if os.path.exists(common_paths):
        return common_paths
    else:
        print(f"The path folder does not exist at {common_paths}.")
        return None

def find_wechatocr_exe():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    wechatocr_path = os.path.join(script_dir, 'path', 'WeChatOCR', 'WeChatOCR.exe')
    if os.path.isfile(wechatocr_path):
        return wechatocr_path
    else:
        print(f"The WeChatOCR.exe does not exist at {wechatocr_path}.")
        return None

def wechat_ocr(image_path):
    wechat_path = find_wechat_path()
    wechatocr_path = find_wechatocr_exe()
    if not wechat_path or not wechatocr_path:
        return []  # 返回空结果
    
    wcocr.init(wechatocr_path, wechat_path)
    result = wcocr.ocr(image_path)
    texts = []

    for temp in result['ocr_response']:
        text = temp['text']
        if isinstance(text, bytes):
            text = text.decode('utf-8', errors='ignore')
        texts.append(text)
    
    return texts

def save_to_docx(texts, output_path):
    doc = Document()

    for text in texts:
        # 添加段落并设置宋体字体
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        run.font.name = '宋体'

        # 设置字体为宋体 (兼容中文设置)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 设置字体大小为五号字体 (10.5 磅)
        run.font.size = Pt(10.5)
    
    doc.save(output_path)

def process_all_images():
    script_dir = os.path.dirname(os.path.abspath(__file__))
    src_folder = os.path.join(script_dir, 'src')
    docx_folder = os.path.join(script_dir, 'docx')

    if not os.path.exists(docx_folder):
        os.makedirs(docx_folder)

    # 支持的图像格式
    image_extensions = ('.jpg', '.jpeg', '.png', '.bmp', '.tif')

    # 遍历 src 文件夹及其所有子文件夹
    for root, dirs, files in os.walk(src_folder):
        for file in files:
            if file.lower().endswith(image_extensions):
                image_path = os.path.join(root, file)
                relative_path = os.path.relpath(root, src_folder)
                docx_folder_path = os.path.join(docx_folder, relative_path)

                # 确保 docx 文件夹路径存在
                if not os.path.exists(docx_folder_path):
                    os.makedirs(docx_folder_path)

                # 处理图片文件
                print(Fore.GREEN + f"正在处理: {os.path.relpath(image_path, script_dir)}" + Style.RESET_ALL)
                texts = wechat_ocr(image_path)
                image_name = os.path.splitext(file)[0]
                output_docx = os.path.join(docx_folder_path, f'{image_name}_OCR.docx')
                save_to_docx(texts, output_docx)
                # 显示相对路径
                relative_docx_path = os.path.relpath(output_docx, script_dir)
                print(f"OCR 结果已保存到： {relative_docx_path}\n")

if __name__ == '__main__':
    init(autoreset=True)  # 初始化 colorama
    process_all_images()
    print(Fore.RED + "全部文件处理完成，请按 Enter 键退出……" + Style.RESET_ALL)
    input()
